"""Template INACAP v5 — Formato del Informe profesional.
Basado en reglas institucionales INACAP (Arial 11pt body, 14pt/12pt headings, single spacing).
Portada validada contra gold standard del usuario.
Uso: importar y llamar funciones. NO copiar este código.
"""
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import os

ASSETS = os.path.expanduser("~/.claude/skills/inacap/assets")

# ── Constantes de formato (colores institucionales INACAP) ──────────────────

# INACAP brand (extraído del SVG oficial: rojo #ed1c24, negro #231f20)
INACAP_RED = "ed1c24"        # Rojo corporativo INACAP
INACAP_DARK = "231f20"       # Negro corporativo INACAP

# Portada
COVER_TITLE = "C00000"       # Rojo título (más oscuro para impresión)
COVER_SUBTITLE = "4C4C4C"    # Gris subtítulo
COVER_META = "404040"        # Gris metadata
COVER_DATE = "606060"        # Gris fecha

# Headings
HEADING1_COLOR = "231f20"    # Negro corporativo
HEADING2_COLOR = "ed1c24"    # Rojo INACAP
HEADING3_COLOR = "404040"    # Gris oscuro

# Tablas
TABLE_HEADER_BG = "231f20"   # Fondo cabecera (negro corporativo)
TABLE_HEADER_TEXT = "FFFFFF"  # Texto cabecera
TABLE_BORDER = "BFBFBF"       # Bordes de tabla
ROW_ALT_A = "FFFFFF"          # Fila blanca
ROW_ALT_B = "E7E6E6"          # Fila gris alternada

# Acentos
DIVIDER_RED = "ed1c24"       # Líneas divisorias
FOOTER_TEXT = "666666"        # Texto footer
FIGURE_GRAY = "595959"        # Pie de figura

# Medidas de página (v2: 2.50 cm, no 2.54 cm)
MARGIN_CM = 2.50
HEADER_DIST_CM = 0.79
FOOTER_DIST_CM = 0.96


# ── API Pública ───────────────────────────────────────────────────────────────

def create_doc(carrera, asignatura, titulo, profesor, integrantes, fecha, sede,
               subtitulo=""):
    """Crea documento con portada: banner wide en header, metadata centrada sin labels.

    Arquitectura 2 secciones:
      Sección 1 → Portada (header con banner wide, sin footer)
      Sección 2 → Cuerpo (header con logo + footer con página)

    Retorna doc.
    """
    _verificar_assets()

    doc = Document()

    # ── Sección 1: Portada (header con banner wide, sin footer) ──
    s1 = doc.sections[0]
    _setup_page(s1)
    _enable_title_page(s1)
    s1.first_page_header.is_linked_to_previous = False
    _add_cover_banner_header(s1.first_page_header)

    # Espaciado vertical
    for _ in range(12):
        _empty(doc)

    _center(doc, titulo, Pt(16), bold=True, color=COVER_TITLE)
    if subtitulo:
        _center(doc, subtitulo, Pt(14), color=COVER_SUBTITLE)

    _empty(doc)  # espacio

    _center(doc, profesor, Pt(12), color=COVER_META)

    _empty(doc)  # espacio

    for nombre in integrantes:
        _center(doc, nombre, Pt(12), color=COVER_META)

    _empty(doc)  # espacio

    _center(doc, fecha, Pt(12), color=COVER_DATE)

    _empty(doc)  # espacio

    _center(doc, sede, Pt(12), color=COVER_META)

    _empty(doc)  # espacio final

    # ── Sección 2: Cuerpo (header con logo + footer con página) ──
    s2 = doc.add_section()
    _setup_page(s2)
    _remove_title_page(s2)
    s2.header.is_linked_to_previous = False
    _add_header_logo(s2.header)
    s2.footer.is_linked_to_previous = False
    _setup_footer(s2.footer, asignatura)

    # ── Estilos ──
    _setup_styles(doc)

    return doc


def heading(doc, text, level=1):
    """Agrega heading. level 1-6 según jerarquía. H1 lleva línea divisoria."""
    h = doc.add_heading(text, level=level)
    if level == 1:
        pPr = h._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="{DIVIDER_RED}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    return h


def body(doc, text):
    """Párrafo Normal justificado Arial 11pt (regla institucional INACAP)."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return p


def section(doc, title, paragraphs, level=1):
    """Heading + párrafos de body."""
    heading(doc, title, level=level)
    for text in paragraphs:
        body(doc, text)


def bullet(doc, text):
    """Bullet point: '•  ' + texto, 11pt, line_spacing=1.0."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(f"•  {text}")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def bullets(doc, items):
    """Lista de bullet points."""
    for item in items:
        bullet(doc, item)


def sub_heading(doc, text, color="red"):
    """Sub-heading coloreado: rojo INACAP #ed1c24 o negro #231f20, 11pt bold."""
    colors = {
        "red": RGBColor(0xED, 0x1C, 0x24),
        "dark": RGBColor(0x23, 0x1F, 0x20),
    }
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = colors.get(color, colors["red"])
    run.font.name = "Arial"
    return p


def table(doc, headers, rows, col_widths=None):
    """Tabla formato v2: 6 bordes #AAAAAA, header #1F4E79/blanco, filas alternadas.

    col_widths: lista opcional de anchos en Cm, ej. [Cm(4), Cm(10)].
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    # ── Header row ──
    for i, header_text in enumerate(headers):
        cell = t.rows[0].cells[i]
        _set_cell_borders(cell, TABLE_BORDER)  # borde gris para header
        _shade_cell(cell, TABLE_HEADER_BG)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Data rows ──
    for ri, row_data in enumerate(rows):
        bg = ROW_ALT_A if ri % 2 == 0 else ROW_ALT_B
        for ci, cell_text in enumerate(row_data):
            cell = t.rows[ri + 1].cells[ci]
            _set_cell_borders(cell, TABLE_BORDER)
            _shade_cell(cell, bg)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(11)
            run.font.name = "Arial"

    # ── Anchos de columna ──
    if col_widths:
        for row in t.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    doc.add_paragraph()  # espacio post-tabla
    return t


def page_break(doc):
    """Salto de página."""
    doc.add_page_break()


def figure(doc, number, description):
    """Pie de figura: 9pt italic gris #595959 centrado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(4)
    run = p.add_run(f"Figura {number}: {description}")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    run.font.name = "Arial"
    return p


def insert_toc(doc, levels=3):
    """Inserta campo de Tabla de Contenido. Usar justo despues de create_doc().

    El usuario debe abrir el DOCX en Word, seleccionar todo (Ctrl+A)
    y presionar F9 para generar la tabla de contenido.
    """
    # Titulo "Tabla de Contenido" con estilo H1
    heading(doc, "Tabla de Contenido", level=1)

    # Campo TOC
    p = doc.add_paragraph()
    run = p.add_run()

    fldChar_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fldChar_begin)

    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-{levels}" \\h \\z \\u '
        f'</w:instrText>'
    )
    run._r.append(instrText)

    fldChar_separate = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
    )
    run._r.append(fldChar_separate)

    run2 = p.add_run("[Actualizar tabla de contenido: Ctrl+A, luego F9]")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.font.name = "Arial"

    fldChar_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run2._r.append(fldChar_end)

    doc.add_page_break()
    return p


def image(doc, path, width_cm=None, description=None, number=None):
    """Inserta imagen centrada en el cuerpo del documento.

    Si se proporcionan description y number, agrega pie de figura
    automaticamente con figure().
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    if width_cm:
        run.add_picture(path, width=Cm(width_cm))
    else:
        run.add_picture(path)
    if description is not None and number is not None:
        figure(doc, number, description)
    return p


def cite_apa(authors, year, page=None):
    """Cita APA 7ma parentetica. Retorna string.

    Usa "&" entre autores (regla APA 7 para citas parenteticas).
    1 autor   → "(Apellido, Año)"
    2 autores → "(Apellido1 & Apellido2, Año)"
    3+ autores → "(Apellido1 et al., Año)"
    page: numero de pagina opcional → "(Apellido, Año, p. 42)"
    """
    if isinstance(authors, str):
        authors = [authors]

    n = len(authors)
    if n == 1:
        names = authors[0]
    elif n == 2:
        names = f"{authors[0]} & {authors[1]}"
    else:
        names = f"{authors[0]} et al."

    if page:
        return f"({names}, {year}, p. {page})"
    return f"({names}, {year})"


def cite_apa_narrative(authors, year, page=None):
    """Cita APA 7ma narrativa (autor como parte de la oracion). Retorna string.

    Usa "y" entre autores (regla APA 7 para citas narrativas en espanol).
    1 autor   → "Apellido (Año)"
    2 autores → "Apellido1 y Apellido2 (Año)"
    3+ autores → "Apellido1 et al. (Año)"
    page: numero de pagina opcional → "Apellido (Año, p. 42)"
    """
    if isinstance(authors, str):
        authors = [authors]

    n = len(authors)
    if n == 1:
        names = authors[0]
    elif n == 2:
        names = f"{authors[0]} y {authors[1]}"
    else:
        names = f"{authors[0]} et al."

    if page:
        return f"{names} ({year}, p. {page})"
    return f"{names} ({year})"


def reference_book(doc, author, year, title, publisher, edition=None):
    """Agrega referencia APA 7ma: libro.

    Formato: Apellido, I. (Año). Titulo (X ed.). Editorial.
    """
    ed_text = f" ({edition} ed.)" if edition else ""
    text = f"{author} ({year}). {title}{ed_text}. {publisher}."
    return _reference_paragraph(doc, text)


def reference_book_chapter(doc, author, year, chapter_title, editor, book_title,
                          pages, publisher, edition=None):
    """Agrega referencia APA 7ma: capitulo de libro editado.

    Formato: Apellido, I. (Año). Titulo del capitulo. En I. Editor (Ed.),
             Titulo del libro (X ed., pp. xx-xx). Editorial.
    """
    ed_text = f" ({edition} ed.)" if edition else ""
    text = (f"{author} ({year}). {chapter_title}. En {editor} (Ed.), "
            f"{book_title}{ed_text} (pp. {pages}). {publisher}.")
    return _reference_paragraph(doc, text)


def reference_article(doc, author, year, title, journal, volume, issue=None,
                      pages=None, doi=None):
    """Agrega referencia APA 7ma: articulo de revista.

    Formato: Apellido, I. (Año). Titulo. Revista, Volumen(Issue), paginas. DOI
    """
    text = f"{author} ({year}). {title}. {journal}, {volume}"
    if issue:
        text += f"({issue})"
    if pages:
        text += f", {pages}"
    text += "."
    if doi:
        text += f" https://doi.org/{doi}"
    return _reference_paragraph(doc, text)


def reference_website(doc, author, year, title, site, url):
    """Agrega referencia APA 7ma: pagina web.

    Formato: Apellido, I. (Año). Titulo. Nombre del Sitio. URL
    """
    text = f"{author} ({year}). {title}. {site}. {url}"
    return _reference_paragraph(doc, text)


def save(doc, filepath):
    """Guarda documento creando directorios si es necesario."""
    os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
    doc.save(filepath)


# ── Funciones internas ────────────────────────────────────────────────────────

def _verificar_assets():
    """Verifica que los assets requeridos existan. Falla loud si falta alguno."""
    requeridos = ["inacap_logo.png", "cover_banner_wide.png"]
    for fname in requeridos:
        fpath = os.path.join(ASSETS, fname)
        if not os.path.exists(fpath):
            raise FileNotFoundError(f"Asset faltante: {fpath}")


def _setup_page(section):
    """Página Letter, márgenes 2.50 cm, header/footer distancias v2."""
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(HEADER_DIST_CM)
    section.footer_distance = Cm(FOOTER_DIST_CM)


def _enable_title_page(section):
    """Activa different first page header (titlePg) en la sección."""
    sectPr = section._sectPr
    if sectPr.find(qn('w:titlePg')) is None:
        titlePg = parse_xml(f'<w:titlePg {nsdecls("w")}/>')
        sectPr.append(titlePg)


def _remove_title_page(section):
    """Elimina titlePg de una sección (evita herencia no deseada de sección anterior)."""
    sectPr = section._sectPr
    for el in sectPr.findall(qn('w:titlePg')):
        sectPr.remove(el)



def _add_header_logo(header):
    """Logo INACAP en el header con línea divisoria roja. 8.31cm = tamaño exacto del DOCX de referencia."""
    logo_path = os.path.join(ASSETS, "inacap_logo.png")
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="{DIVIDER_RED}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = hp.add_run()
    run.add_picture(logo_path, width=Cm(8.31))


def _add_cover_banner_header(header):
    """Banner wide en el header de portada. Sin borde, 22.46cm ancho."""
    banner_path = os.path.join(ASSETS, "cover_banner_wide.png")
    hp = header.paragraphs[0]
    run = hp.add_run()
    run.add_picture(banner_path, width=Cm(22.46))



def _setup_footer(footer, asignatura):
    """Footer: asignatura (izq.) + número de página (der.) con tab stop.

    Ref: Node.js template_complete.js — un solo párrafo con tab stop derecho + PAGE field.
    """
    p = footer.paragraphs[0]
    pPr = p._p.get_or_add_pPr()

    # Borde superior gris
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{TABLE_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Tab stop derecho para el número de página
    tabs = parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        f'<w:tab w:val="right" w:pos="9360"/>'
        f'</w:tabs>'
    )
    pPr.append(tabs)

    # Texto: "Asignatura." + tab + PAGE field
    r1 = p.add_run(f"{asignatura}.")
    r1.font.size = Pt(10)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Tab character
    tab_run = p.add_run()
    tab_run._r.append(parse_xml(f'<w:tab {nsdecls("w")}/>'))

    # PAGE field
    r2 = p.add_run()
    r2.font.size = Pt(10)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r2._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2._r.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    ))
    r2._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))



def _setup_styles(doc):
    """Configura estilos según valores exactos extraídos del DOCX v2 de referencia."""

    # ── Normal: 11pt, single spacing (regla institucional INACAP) ──
    sn = doc.styles["Normal"]
    sn.font.name = "Arial"
    sn.font.size = Pt(11)
    sn.font.color.rgb = RGBColor(0, 0, 0)
    sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sn.paragraph_format.space_after = Pt(0)
    sn.paragraph_format.line_spacing = 1.0

    # ── Heading 1: 14pt bold (regla institucional) ──
    s1 = doc.styles["Heading 1"]
    s1.font.name = "Arial"
    s1.font.size = Pt(14)
    s1.font.bold = True
    s1.font.color.rgb = RGBColor(0x23, 0x1F, 0x20)
    s1.paragraph_format.space_before = Pt(12)
    s1.paragraph_format.space_after = Pt(4)
    s1.paragraph_format.line_spacing = 1.0
    _set_keep_next(s1)
    _set_keep_lines(s1)

    # ── Heading 2: 12pt bold (regla institucional) ──
    s2 = doc.styles["Heading 2"]
    s2.font.name = "Arial"
    s2.font.size = Pt(12)
    s2.font.bold = True
    s2.font.color.rgb = RGBColor(0xED, 0x1C, 0x24)
    s2.paragraph_format.space_before = Pt(12)
    s2.paragraph_format.space_after = Pt(4)
    s2.paragraph_format.line_spacing = 1.0
    _set_keep_next(s2)
    _set_keep_lines(s2)

    # ── Heading 3: 12pt bold, mismo tamaño que H2 ──
    s3 = doc.styles["Heading 3"]
    s3.font.name = "Arial"
    s3.font.size = Pt(12)
    s3.font.bold = True
    s3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    s3.paragraph_format.space_before = Pt(10)
    s3.paragraph_format.space_after = Pt(4)
    s3.paragraph_format.line_spacing = 1.0
    _set_keep_next(s3)
    _set_keep_lines(s3)

    # ── Heading 4: 11pt bold ──
    s4 = doc.styles["Heading 4"]
    s4.font.name = "Arial"
    s4.font.size = Pt(11)
    s4.font.bold = True
    s4.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    s4.paragraph_format.space_before = Pt(10)
    s4.paragraph_format.space_after = Pt(2)
    s4.paragraph_format.line_spacing = 1.0
    _set_keep_next(s4)
    _set_keep_lines(s4)

    # ── Heading 5: 11pt bold ──
    s5 = doc.styles["Heading 5"]
    s5.font.name = "Arial"
    s5.font.size = Pt(11)
    s5.font.bold = True
    s5.font.color.rgb = RGBColor(0, 0, 0)
    s5.paragraph_format.space_before = Pt(8)
    s5.paragraph_format.space_after = Pt(2)
    s5.paragraph_format.line_spacing = 1.0
    _set_keep_next(s5)
    _set_keep_lines(s5)

    # ── Heading 6: 11pt bold ──
    s6 = doc.styles["Heading 6"]
    s6.font.name = "Arial"
    s6.font.size = Pt(11)
    s6.font.bold = True
    s6.font.color.rgb = RGBColor(0, 0, 0)
    s6.paragraph_format.space_before = Pt(8)
    s6.paragraph_format.space_after = Pt(2)
    s6.paragraph_format.line_spacing = 1.0
    _set_keep_next(s6)
    _set_keep_lines(s6)


def _set_keep_next(style):
    """Establece keepNext en el XML del estilo."""
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        style.element.append(pPr)
    if pPr.find(qn('w:keepNext')) is None:
        keep_next = parse_xml(f'<w:keepNext {nsdecls("w")}/>')
        pPr.append(keep_next)


def _set_keep_lines(style):
    """Establece keepLines en el XML del estilo."""
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        style.element.append(pPr)
    if pPr.find(qn('w:keepLines')) is None:
        keep_lines = parse_xml(f'<w:keepLines {nsdecls("w")}/>')
        pPr.append(keep_lines)


def _center(doc, text, size, bold=False, color=None, space_before=None, space_after=None):
    """Texto centrado Arial con spacing controlado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    else:
        p.paragraph_format.space_before = Pt(0)
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    else:
        p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.font.name = "Arial"
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return p


def _empty(doc):
    """Párrafo vacío sin spacing para espaciado controlado en portada."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def _divider_line(doc):
    """Línea divisoria roja para la portada."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{DIVIDER_RED}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _shade_cell(cell, color_hex):
    """Aplica sombreado de fondo a una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tcPr.append(shading)


def _set_cell_borders(cell, color_hex, sz="4"):
    """Aplica los 6 bordes a una celda (top, left, bottom, right, insideH, insideV)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcBorders = parse_xml(borders_xml)
    tcPr.append(tcBorders)


def _set_table_borders(table, color_hex, sz="4"):
    """Aplica bordes a nivel de tabla (insideH e insideV)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tblBorders>'
    )
    tblBorders = parse_xml(borders_xml)
    tblPr.append(tblBorders)


def _reference_paragraph(doc, text):
    """Parrafo con sangria francesa APA 7ma (1.27 cm)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    # Sangria francesa via XML: left=720 twips (~1.27cm), hanging=720 twips
    pPr = p._p.get_or_add_pPr()
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="720" w:hanging="720"/>'
    )
    pPr.append(ind)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Arial"
    return p
